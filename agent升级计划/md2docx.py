"""Convert eBay listing markdown to formatted Word document."""
import sys, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_md(md_path):
    with open(md_path, encoding="utf-8") as f:
        return f.read()

def build_docx(content, output_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # H1 heading
        if line.startswith("# ") and not line.startswith("## "):
            h = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # H2 heading / separator
        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            doc.add_paragraph("")
            i += 1
            continue

        # Table - detect pipe-delimited rows
        if line.startswith("|") and line.endswith("|") and (i+1 < len(lines) and lines[i+1].startswith("|") and "---" in lines[i+1]):
            # Collect all consecutive table rows
            tbl_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].split("|")[1:-1]]
                tbl_rows.append(row)
                i += 1
            # Remove separator row (second row with ---)
            header = tbl_rows[0]
            data = [r for r in tbl_rows[1:] if not all(re.match(r"^[-:\s]+$", c) for c in r)]
            if data:
                table = doc.add_table(rows=1+len(data), cols=len(header))
                table.style = 'Light Grid Accent 1'
                for ci, h in enumerate(header):
                    table.rows[0].cells[ci].text = h
                for ri, row in enumerate(data):
                    for ci, val in enumerate(row):
                        if ci < len(header):
                            table.rows[ri+1].cells[ci].text = val
            continue

        # Bold bullet points (**KEY** - text)
        m = re.match(r"^\*\*(.+?)\*\*\s*[-–]\s*(.+)", line)
        if m:
            p = doc.add_paragraph()
            run_b = p.add_run(m.group(1))
            run_b.bold = True
            p.add_run(" — " + m.group(2))
            i += 1
            continue

        # Bullet list items
        if re.match(r"^- ", line):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
            continue

        # Code blocks (```...```)
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            code_text = "\n".join(code_lines).strip()
            if code_text:
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            continue

        # Regular paragraph
        if line.strip():
            doc.add_paragraph(line.strip())

        i += 1

    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == "__main__":
    md_path = sys.argv[1] if len(sys.argv) > 1 else "listing_output.md"
    out_path = sys.argv[2] if len(sys.argv) > 2 else md_path.replace(".md", ".docx")
    content = parse_md(md_path)
    build_docx(content, out_path)
